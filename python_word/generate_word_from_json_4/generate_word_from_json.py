import json
from docx import Document
from docx.shared import RGBColor

def generate_word_from_json(json_path, output_word_path):
    """
    根據 JSON 檔案內容產生 Word 文件，處理段落的細節。
    :param json_path: 輸入的 JSON 檔案路徑
    :param output_word_path: 輸出的 Word 檔案路徑
    """
    try:
        # 讀取 JSON 檔案
        with open(json_path, "r", encoding="utf-8") as f:
            data = json.load(f)

        # 創建一個新的 Word 文件
        doc = Document()

        # 處理段落
        for paragraph_data in data.get("paragraphs", []):
            para = doc.add_paragraph()
            para.text = paragraph_data["text"]
            
            # 檢查段落樣式類型
            if paragraph_data["style"]:  # 如果有樣式
                try:
                    # 應用段落樣式
                    para.style = paragraph_data["style"]
                except Exception as e:
                    print(f"警告: 無法應用段落樣式 {paragraph_data['style']}: {e}")

            # 處理段落內的 run
            for run_data in paragraph_data.get("runs", []):
                run = para.add_run(run_data["text"])

                # 應用 run 的樣式屬性
                if run_data["bold"]:
                    run.bold = True
                if run_data["italic"]:
                    run.italic = True
                if run_data["underline"]:
                    run.underline = True
                if run_data["font_name"]:
                    run.font.name = run_data["font_name"]
                if run_data["font_size"]:
                    run.font.size = run_data["font_size"]
                
                # 處理字體顏色，如果顏色存在且為列表，轉換為 RGBColor 物件
                if run_data["font_color"]:
                    color = run_data["font_color"]
                    if isinstance(color, list) and len(color) == 3:
                        run.font.color.rgb = RGBColor(*color)

        # 處理表格
        for table_data in data.get("tables", []):
            table = doc.add_table(rows=len(table_data["rows"]), cols=len(table_data["rows"][0]))

            # 填入表格內容
            for row_index, row_data in enumerate(table_data["rows"]):
                for cell_index, cell_data in enumerate(row_data):
                    cell = table.cell(row_index, cell_index)
                    cell.text = cell_data["text"]
                    
                    # 應用儲存格樣式
                    if cell_data["style"]:
                        try:
                            # 儲存格樣式應用
                            cell.paragraphs[0].style = cell_data["style"]
                        except Exception as e:
                            print(f"警告: 無法應用儲存格樣式 {cell_data['style']}: {e}")
                    
                    # 處理儲存格內的 run
                    for run_data in cell_data.get("runs", []):
                        run = cell.paragraphs[0].add_run(run_data["text"])

                        # 應用 run 的樣式屬性
                        if run_data["bold"]:
                            run.bold = True
                        if run_data["italic"]:
                            run.italic = True
                        if run_data["underline"]:
                            run.underline = True
                        if run_data["font_name"]:
                            run.font.name = run_data["font_name"]
                        if run_data["font_size"]:
                            run.font.size = run_data["font_size"]
                        
                        # 處理字體顏色
                        if run_data["font_color"]:
                            color = run_data["font_color"]
                            if isinstance(color, list) and len(color) == 3:
                                run.font.color.rgb = RGBColor(*color)

        # 儲存 Word 文件
        doc.save(output_word_path)
        print(f"成功生成 Word 檔案：{output_word_path}")

    except Exception as e:
        print(f"處理 JSON 檔案或生成 Word 檔案時發生錯誤：{e}")

if __name__ == "__main__":
    json_file = "detailed_output.json"  # 輸入的 JSON 檔案路徑
    word_file = "generated.docx"  # 輸出的 Word 檔案路徑
    generate_word_from_json(json_file, word_file)
