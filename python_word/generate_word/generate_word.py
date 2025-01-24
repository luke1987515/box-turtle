import datetime
from docx import Document

# 建立 Word 文件
doc = Document()
doc.add_heading('自動化產生的 Word 文件', level=1)
doc.add_paragraph('這是第一段文字內容。')
doc.add_paragraph('這是第二段，還可以加入樣式。', style='Intense Quote')

# 動態產生檔名（加上日期時間）
timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
filename = f"自動化文件_{timestamp}.docx"

# 儲存文件
doc.save(filename)
print(f"文件已儲存為：{filename}")
