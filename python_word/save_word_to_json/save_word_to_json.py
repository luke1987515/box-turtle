import json
from docx import Document

def save_word_data_to_json(file_path, output_path):
    """
    讀取 Word 文件內容，並將內容、樣式與位置資訊存為 JSON 檔案。
    :param file_path: Word 文件路徑
    :param output_path: 輸出 JSON 文件路徑
    """
    try:
        doc = Document(file_path)
        data = {"paragraphs": [], "tables": []}

        # 段落內容
        for para_index, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():
                data["paragraphs"].append({
                    "index": para_index,
                    "text": paragraph.text,
                    "style": paragraph.style.name
                })

        # 表格內容
        for table_index, table in enumerate(doc.tables):
            table_data = {"table_index": table_index, "rows": []}
            for row_index, row in enumerate(table.rows):
                row_data = []
                for cell_index, cell in enumerate(row.cells):
                    if cell.text.strip():
                        row_data.append({
                            "cell_index": cell_index,
                            "text": cell.text.strip(),
                            "style": cell.paragraphs[0].style.name
                        })
                table_data["rows"].append(row_data)
            data["tables"].append(table_data)

        # 儲存為 JSON 文件
        with open(output_path, "w", encoding="utf-8") as f:
            json.dump(data, f, indent=4, ensure_ascii=False)
        print(f"內容已儲存為 JSON 檔案：{output_path}")

    except Exception as e:
        print(f"處理文件時發生錯誤：{e}")

if __name__ == "__main__":
    input_file = "sample.docx"  # 替換成你的 Word 文件路徑
    output_file = "output.json"  # 輸出的 JSON 文件路徑
    save_word_data_to_json(input_file, output_file)
