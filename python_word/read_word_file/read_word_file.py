from docx import Document

def read_word_file(file_path):
    """
    讀取 .docx 文件內容並顯示
    :param file_path: Word 檔案的路徑
    """
    try:
        # 載入 Word 文件
        doc = Document(file_path)

        # 讀取段落內容
        print("段落內容：")
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # 過濾空白段落
                print(paragraph.text)

        # 讀取表格內容
        if doc.tables:
            print("\n表格內容：")
            for table in doc.tables:
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    print("\t".join(row_data))  # 用 Tab 分隔列內容

    except Exception as e:
        print(f"讀取檔案時發生錯誤：{e}")

if __name__ == "__main__":
    # 替換為你的 Word 檔案名稱（需包含路徑）
    file_path = "sample.docx"  # 預設檔案名稱為 sample.docx
    read_word_file(file_path)
