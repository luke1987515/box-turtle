import json
from docx import Document

def generate_word_from_json(json_path, output_word_path):
    """
    根據 JSON 檔案內容產生 Word 文件，處理段落的細節。
    :param json_path: 輸入的 JSON 檔案路徑
    :param output_word_path: 輸出的 Word 檔案路徑
    """
    try:
        # 讀取 JSON 檔案
        with open(json_path, "r", encoding="utf-8") as f:
            data = json.load(f)

        # 創建一個新的 Word 文件
        doc = Document()

        # 處理段落
        for paragraph in data.get("paragraphs", []):
            para = doc.add_paragraph()
            # 根據段落中的每個 run 來設置文字樣式
            para.text = paragraph["text"]
            para.style = paragraph["style"]

            # 如果有段落內的文字運行 (run) 資料
            # 可以對每個 run 進行樣式設置
            # 例如，你可以把一部分文字設定為加粗、斜體等
            for run in para.runs:
                run.style = paragraph["style"]  # 設定樣式

        # 處理表格
        for table in data.get("tables", []):
            table_data = table.get("rows", [])
            if table_data:
                # 創建表格
                table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))

                # 填入表格內容
                for row_index, row in enumerate(table_data):
                    for cell_index, cell in enumerate(row):
                        # 填入儲存格內容
                        table.cell(row_index, cell_index).text = cell["text"]
                        # 設定儲存格樣式
                        table.cell(row_index, cell_index).paragraphs[0].style = cell["style"]

        # 儲存 Word 文件
        doc.save(output_word_path)
        print(f"成功生成 Word 檔案：{output_word_path}")

    except Exception as e:
        print(f"處理 JSON 檔案或生成 Word 檔案時發生錯誤：{e}")

if __name__ == "__main__":
    # 設定 JSON 檔案與輸出 Word 檔案路徑
    json_file = "output.json"  # 輸入的 JSON 檔案路徑
    word_file = "generated.docx"  # 輸出的 Word 檔案路徑
    generate_word_from_json(json_file, word_file)
