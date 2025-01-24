from docx import Document

def read_word_with_styles_and_positions(file_path):
    """
    讀取 .docx 文件，顯示內容、樣式與位置資訊。
    :param file_path: Word 檔案的路徑
    """
    try:
        # 載入 Word 文件
        doc = Document(file_path)

        # 讀取段落內容與位置
        print("段落內容、樣式與位置：")
        for para_index, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():  # 過濾空白段落
                print(f"段落索引: {para_index}")
                print(f"段落文字: {paragraph.text}")
                print(f"段落樣式: {paragraph.style.name}\n")

        # 讀取段落內文字內容與位置
        print("段落內文字與樣式位置：")
        for para_index, paragraph in enumerate(doc.paragraphs):
            for run_index, run in enumerate(paragraph.runs):
                if run.text.strip():  # 過濾空白文字
                    print(f"段落索引: {para_index}, 文字索引: {run_index}")
                    print(f"文字: {run.text}")
                    print(f"樣式: {run.style.name}\n")

        # 讀取表格內容與位置
        if doc.tables:
            print("表格內容、樣式與位置：")
            for table_index, table in enumerate(doc.tables):
                for row_index, row in enumerate(table.rows):
                    for cell_index, cell in enumerate(row.cells):
                        cell_text = cell.text.strip()
                        if cell_text:
                            print(f"表格索引: {table_index}, 列索引: {row_index}, 儲存格索引: {cell_index}")
                            print(f"儲存格文字: {cell_text}")
                            print(f"儲存格樣式: {cell.paragraphs[0].style.name}\n")  # 使用儲存格第一段的樣式

    except Exception as e:
        print(f"讀取檔案時發生錯誤：{e}")

if __name__ == "__main__":
    # 替換為你的 Word 檔案名稱（需包含路徑）
    file_path = "sample.docx"  # 預設檔案名稱為 sample.docx
    read_word_with_styles_and_positions(file_path)
