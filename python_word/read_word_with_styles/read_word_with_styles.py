from docx import Document

def read_word_with_styles(file_path):
    """
    讀取 .docx 文件的內容和樣式，並顯示段落及文字的樣式。
    :param file_path: Word 檔案的路徑
    """
    try:
        # 載入 Word 文件
        doc = Document(file_path)

        # 讀取段落內容和樣式
        print("段落內容與樣式：")
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # 過濾空白段落
                print(f"段落文字: {paragraph.text}")
                print(f"段落樣式: {paragraph.style.name}\n")  # 顯示段落樣式名稱

        # 讀取表格內容（包含樣式）
        if doc.tables:
            print("\n表格內容與樣式：")
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            print(f"儲存格文字: {cell_text}")
                            print(f"儲存格樣式: {cell.paragraphs[0].style.name}\n")  # 假設儲存格的第一段樣式
        
        # 讀取每段落中的文字樣式（若有）
        print("\n段落文字內的字詞樣式：")
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if run.text.strip():  # 過濾空白文字
                    print(f"文字: {run.text}")
                    print(f"樣式: {run.style.name}\n")  # 顯示文字的樣式名稱

    except Exception as e:
        print(f"讀取檔案時發生錯誤：{e}")

if __name__ == "__main__":
    # 替換為你的 Word 檔案名稱（需包含路徑）
    file_path = "sample.docx"  # 預設檔案名稱為 sample.docx
    read_word_with_styles(file_path)
