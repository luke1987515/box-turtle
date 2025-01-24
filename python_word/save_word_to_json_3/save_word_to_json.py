import json
from docx import Document

def save_word_data_to_json(file_path, output_path):
    """
    讀取 Word 文件內容，並將內容、樣式與位置資訊存為 JSON 檔案。
    :param file_path: Word 文件路徑
    :param output_path: 輸出 JSON 文件路徑
    """
    try:
        doc = Document(file_path)
        data = {"paragraphs": [], "tables": []}

        # 檢索段落
        for para_index, paragraph in enumerate(doc.paragraphs):
            para_data = {
                "index": para_index,
                "text": paragraph.text,
                "style": paragraph.style.name if paragraph.style else None,  # 儲存段落樣式
                "runs": []
            }
            
            # 檢索段落中的 run
            for run_index, run in enumerate(paragraph.runs):
                run_data = {
                    "run_index": run_index,
                    "text": run.text,
                    "bold": run.bold,
                    "italic": run.italic,
                    "underline": run.underline,
                    "font_name": run.font.name,
                    "font_size": run.font.size,
                    "font_color": run.font.color.rgb if run.font.color else None  # 儲存顏色資訊
                }
                para_data["runs"].append(run_data)

            # 把段落資料加入到資料中
            data["paragraphs"].append(para_data)

        # 檢索表格
        for table_index, table in enumerate(doc.tables):
            table_data = {"table_index": table_index, "rows": []}

            for row_index, row in enumerate(table.rows):
                row_data = []
                for cell_index, cell in enumerate(row.cells):
                    cell_data = {
                        "cell_index": cell_index,
                        "text": cell.text.strip(),
                        "style": cell.paragraphs[0].style.name if cell.paragraphs else None,  # 儲存儲存格的段落樣式
                        "runs": []
                    }

                    # 檢索單元格中的 run
                    for run_index, run in enumerate(cell.paragraphs[0].runs):
                        run_data = {
                            "run_index": run_index,
                            "text": run.text,
                            "bold": run.bold,
                            "italic": run.italic,
                            "underline": run.underline,
                            "font_name": run.font.name,
                            "font_size": run.font.size,
                            "font_color": run.font.color.rgb if run.font.color else None
                        }
                        cell_data["runs"].append(run_data)

                    row_data.append(cell_data)
                table_data["rows"].append(row_data)

            data["tables"].append(table_data)

        # 儲存為 JSON 文件
        with open(output_path, "w", encoding="utf-8") as f:
            json.dump(data, f, indent=4, ensure_ascii=False)
        print(f"內容已儲存為 JSON 檔案：{output_path}")

    except Exception as e:
        print(f"處理文件時發生錯誤：{e}")

if __name__ == "__main__":
    input_file = "sample.docx"  # 替換成你的 Word 文件路徑
    output_file = "detailed_output.json"  # 輸出的 JSON 文件路徑
    save_word_data_to_json(input_file, output_file)
