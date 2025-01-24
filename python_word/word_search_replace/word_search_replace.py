from docx import Document

def search_and_replace(doc_path, keyword, replacement, output_path):
    # 打開 Word 文件
    doc = Document(doc_path)
    
    # 遍歷每一段落
    for paragraph in doc.paragraphs:
        if keyword in paragraph.text:
            paragraph.text = paragraph.text.replace(keyword, replacement)
    
    # 遍歷表格中的文字（如果有表格的話）
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if keyword in cell.text:
                    cell.text = cell.text.replace(keyword, replacement)
    
    # 儲存修改後的文件
    doc.save(output_path)
    print(f"已完成替換並儲存到 {output_path}")

# 測試範例
# doc_path = "example.docx"  # 原始 Word 文件路徑
doc_path = "AIC Test Report_Template(直式)_2021_Functionality_Test_JBOD_1206.docx"  # 原始 Word 文件路徑
# keyword = "關鍵字"          # 要搜尋的關鍵字
keyword = "XXXXXXXXXXX"          # 要搜尋的關鍵字
# replacement = "替換文字"     # 替換成的文字
replacement = "X-X-X-X-X-X"     # 替換成的文字
output_path = "output.docx" # 輸出的文件路徑

search_and_replace(doc_path, keyword, replacement, output_path)
